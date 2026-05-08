from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from jinja2 import Environment, FileSystemLoader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.core.application_tracker import append_application_record
from app.core.ats_scorer import render_ats_report, score_resume_ats
from app.core.interview_prep import generate_interview_prep, render_interview_prep
from app.core.missing_skills import analyze_missing_skills, render_missing_skills_report
from app.core.resume_diff import create_resume_diff
from app.models import CandidateProfile, EvidenceMatch, FitScore, JobDescription, Project, ResumeRules
from app.validators.truthfulness_validator import render_truthfulness_report, validate_claims

TEMPLATE_DIR = Path(__file__).resolve().parents[1] / "templates"


def render_resume(job: JobDescription, profile: CandidateProfile, projects: list[Project], fit: FitScore, evidence: list[EvidenceMatch]) -> str:
    template_name = "resumes/genai_qa_resume.md.j2" if "AI Testing" in fit.best_positioning else "resumes/sdet_fullstack_resume.md.j2"
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR), autoescape=False)
    return env.get_template(template_name).render(job=job, profile=profile, projects=projects, fit_score=fit, evidence_matches=[e for e in evidence if e.evidence_strength != "Missing"])


def render_base_resume(profile: CandidateProfile, projects: list[Project]) -> str:
    lines = [f"# {profile.candidate.get('name', 'Candidate Name')}", "", profile.candidate.get("headline", "QA Candidate"), "", "## Professional Summary", "", profile.summary.get("default", ""), "", "## Skills", ""]
    for group, skills in profile.skills.items():
        lines.append(f"- {group}: {', '.join(skills)}")
    lines += ["", "## Projects", ""]
    for project in projects:
        lines += [f"### {project.name}", project.type, ""]
        for bullet in project.safe_resume_bullets[:1]:
            lines.append(f"- {bullet}")
        lines.append("")
    return "\n".join(lines)


def render_cover_letter(job: JobDescription, profile: CandidateProfile, fit: FitScore, evidence: list[EvidenceMatch]) -> str:
    strong = [e for e in evidence if e.evidence_strength == "Strong"][:4]
    project_names = _project_names(strong)
    opening = _cover_opening(job, fit, project_names)
    lines = [
        f"# Cover Letter — {job.role_title} at {job.company}",
        "",
        "Dear Hiring Team,",
        "",
        opening,
        "",
        "The specific evidence I would bring to this role is:",
    ]
    for item in strong:
        lines.append(f"- {item.requirement}: {item.resume_bullet}")
    lines += [
        "",
        _cover_closing(job, fit),
        "",
        "Regards,",
        profile.candidate.get("name", "Candidate"),
    ]
    return "\n".join(lines)


def render_recruiter_message(job: JobDescription, fit: FitScore, evidence: list[EvidenceMatch]) -> str:
    strong = [e for e in evidence if e.evidence_strength == "Strong"][:2]
    proof = "; ".join(f"{e.source_project}: {e.requirement}" for e in strong if e.source_project) or "evidence-backed QA project work"
    return (
        f"Hi [Name], I saw the {job.role_title} role at {job.company}. My strongest fit is {fit.best_positioning}. "
        f"The proof is not generic resume wording: {proof}. I have a tailored resume, evidence map, ATS report, and interview-prep pack ready for manual review."
    )


def render_fit_score(job: JobDescription, fit: FitScore) -> str:
    b = fit.breakdown
    missing = [f"- {s}" for s in fit.missing_skills] or ["- None detected from parsed must-haves."]
    risks = [f"- {r}" for r in fit.risks] or ["- No major risk detected."]
    lines = ["# Fit Score Report", "", f"Role: {job.role_title}", f"Company: {job.company}", f"Final Score: {fit.final_score}/100", f"Decision: {fit.decision}", f"Confidence: {fit.confidence}", "", "## Score Breakdown", "", f"- Must-have skill match: {b.must_have_skill_match}/30", f"- Project evidence match: {b.project_evidence_match}/25", f"- Domain relevance: {b.domain_relevance}/15", f"- Seniority match: {b.seniority_match}/10", f"- Tooling/platform match: {b.tooling_platform_match}/10", f"- Resume positioning strength: {b.resume_positioning_strength}/5", f"- Risk check: {b.risk_check}/5", "", "## Best Positioning", "", fit.best_positioning, "", "## Explanation", "", fit.explanation, "", "## Missing Skills", "", *missing, "", "## Risks", "", *risks]
    return "\n".join(lines)


def render_evidence_map(matches: list[EvidenceMatch]) -> str:
    lines = ["# Evidence Map", ""]
    for m in matches:
        lines += [f"## Requirement: {m.requirement}", "", f"Candidate Evidence: {m.candidate_evidence}", "", f"Resume Bullet: {m.resume_bullet or 'No safe bullet generated.'}", "", f"Evidence Strength: {m.evidence_strength}", "", f"Source Project: {m.source_project or 'Profile / Missing'}", "", "---", ""]
    return "\n".join(lines)


def render_notes(job: JobDescription, fit: FitScore) -> str:
    return "\n".join(["# Application Notes", "", f"Role: {job.role_title}", f"Company: {job.company}", f"Decision: {fit.decision}", f"Best positioning: {fit.best_positioning}", "", "## Manual Approval Checklist", "", "- [ ] Resume reviewed", "- [ ] Cover letter reviewed", "- [ ] Recruiter message reviewed", "- [ ] Fit score accepted", "- [ ] ATS report reviewed", "- [ ] Interview prep reviewed", "- [ ] User manually submitted application", "", "This pack is preparation-only. No auto-apply action has been taken."])


def write_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(_strip_md(line))
    doc.save(output_path)


def write_pdf(markdown_text: str, output_path: Path) -> None:
    pdf = canvas.Canvas(str(output_path), pagesize=letter)
    _, height = letter
    y = height - 50
    for raw in markdown_text.splitlines():
        line = _strip_md(raw)[:110]
        if line:
            pdf.drawString(50, y, line)
        y -= 14
        if y < 50:
            pdf.showPage()
            y = height - 50
    pdf.save()


def write_application_pack(output_dir: str | Path, job: JobDescription, profile: CandidateProfile, projects: list[Project], rules: ResumeRules, fit: FitScore, evidence: list[EvidenceMatch]) -> list[str]:
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    base_resume = render_base_resume(profile, projects)
    resume = render_resume(job, profile, projects, fit, evidence)
    missing_report = render_missing_skills_report(analyze_missing_skills(job, profile, projects))
    ats_report = render_ats_report(score_resume_ats(resume, job, evidence))
    interview_report = render_interview_prep(generate_interview_prep(job, fit, evidence))
    files = {
        "01_raw_job_description.txt": job.raw_text,
        "02_parsed_job_description.json": json.dumps(job.model_dump(), indent=2),
        "03_fit_score.md": render_fit_score(job, fit),
        "04_evidence_map.md": render_evidence_map(evidence),
        "05_tailored_resume.md": resume,
        "08_cover_letter.md": render_cover_letter(job, profile, fit, evidence),
        "09_recruiter_message.md": render_recruiter_message(job, fit, evidence),
        "10_application_notes.md": render_notes(job, fit),
        "11_truthfulness_report.md": render_truthfulness_report(validate_claims(resume, evidence, rules)),
        "12_missing_skills_report.md": missing_report,
        "13_resume_diff.md": create_resume_diff(base_resume, resume),
        "14_ats_score_report.md": ats_report,
        "15_interview_prep.md": interview_report,
    }
    written = []
    for name, content in files.items():
        path = out / name
        path.write_text(content, encoding="utf-8")
        written.append(str(path))
    write_docx(resume, out / "06_tailored_resume.docx")
    write_pdf(resume, out / "07_tailored_resume.pdf")
    written += [str(out / "06_tailored_resume.docx"), str(out / "07_tailored_resume.pdf")]
    append_application_record(out.parent / "application_tracker.csv", job, fit, out)
    return written


def _cover_opening(job: JobDescription, fit: FitScore, project_names: list[str]) -> str:
    project_text = ", ".join(project_names[:3]) if project_names else "my QA automation portfolio"
    return (
        f"I am applying for the {job.role_title} role with a focused angle: {fit.best_positioning}. "
        f"My fit is backed by concrete project evidence from {project_text}, not broad claims."
    )


def _cover_closing(job: JobDescription, fit: FitScore) -> str:
    return (
        f"I would be useful in this {job.role_title} role because I can turn ambiguous requirements into testable risks, "
        f"map those risks to automation or manual validation, and communicate gaps honestly. My positioning for this application is {fit.best_positioning}."
    )


def _project_names(evidence: list[EvidenceMatch]) -> list[str]:
    names = []
    for item in evidence:
        if item.source_project and item.source_project not in names:
            names.append(item.source_project)
    return names


def _strip_md(line: str) -> str:
    return re.sub(r"^#+\s*", "", line).replace("**", "")
